from datetime import datetime
from docx import Document
from docx.shared import Pt

output_path = r"k:\6th Sem\IP LAB\stocksense\StockSense_Mini_Project_Report.docx"

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

doc.add_heading('StockSense - Mini Project Report', 0)
meta = doc.add_paragraph()
meta.add_run('Project: ').bold = True
meta.add_run('StockSense - Stock Portfolio Tracker')
meta.add_run('\nTechnology Stack: ').bold = True
meta.add_run('React.js, Tailwind CSS, Axios, Recharts, React Router, Node.js, Express.js, MongoDB Atlas, JWT, bcryptjs')
meta.add_run('\nReport Date: ').bold = True
meta.add_run(datetime.now().strftime('%d %B %Y'))


doc.add_heading('1. Abstract', level=1)
doc.add_paragraph(
    'StockSense is a full-stack MVC web application for tracking stock portfolios. '
    'It allows users to register/login securely, add and manage holdings, monitor live prices, '
    'view profit/loss analytics, maintain transactions, and interact with an integrated AI chatbot. '
    'The system provides portfolio-level summaries and charts for better investment decision support.'
)


doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Retail investors often maintain stock records manually, leading to inconsistent updates, '
    'limited analytics, and poor visibility into portfolio performance. A centralized, real-time '
    'web platform is required to simplify tracking, improve insights, and support action-oriented decisions.'
)


doc.add_heading('3. Objectives', level=1)
for item in [
    'Implement secure user authentication with JWT and password hashing.',
    'Provide CRUD operations for holdings and transactions.',
    'Auto-update holding prices from live market data APIs.',
    'Compute portfolio metrics: total invested, current value, total P&L, and P&L%.',
    'Visualize portfolio data using charts for quick interpretation.',
    'Offer AI chatbot support for stock- and app-related assistance.'
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('4. System Architecture', level=1)
doc.add_paragraph('StockSense follows a layered MVC architecture:')
for item in [
    'Frontend (React): UI, routing, state/context management, charts, AJAX requests.',
    'Backend (Express MVC): route handling, controllers, business logic, authentication middleware.',
    'Database (MongoDB Atlas): persistent storage for users, portfolios, holdings, and transactions.',
    'External APIs: market quote endpoints for real-time price updates and chatbot provider API.'
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('5. Database Design', level=1)
doc.add_paragraph('Main collections and key fields:')
for item in [
    'User: name, email, password, createdAt',
    'Portfolio: userId, name, createdAt',
    'Holding: portfolioId, symbol, companyName, quantity, buyPrice, buyDate, currentPrice, virtuals (pnl, pnlPercent, invested, currentValue)',
    'Transaction: portfolioId, symbol, type (BUY/SELL), quantity, price, date, notes, createdAt'
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('6. Major Modules Implemented', level=1)
sections = {
    'Authentication Module': [
        'User registration and login with JWT token generation (7-day expiry).',
        'Password hashing with bcryptjs.',
        'Protected routes using Bearer token middleware.',
        'Default portfolio creation during user registration.'
    ],
    'Portfolio & Holdings Module': [
        'Add, view, update, and delete holdings.',
        'Automatic BUY transaction logging while adding holdings.',
        'Live quote fetch for symbol lookup and auto-fill while adding stock.',
        'Auto-refresh current prices for all holdings in a portfolio.'
    ],
    'Transaction Module': [
        'Add and list BUY/SELL transactions sorted by date.',
        'SELL logic validates available quantity before execution.',
        'Holdings are adjusted based on BUY/SELL activity.'
    ],
    'Visualization Module': [
        'Portfolio P&L chart (invested vs current value).',
        'Live market chart for current prices by symbol.',
        'Dashboard summary cards for key metrics.'
    ],
    'Chatbot Module': [
        'Protected chat endpoint in backend.',
        'Floating chat widget in frontend.',
        'Provider-based conversational assistance through API key configuration.'
    ]
}
for title, bullets in sections.items():
    doc.add_paragraph(title).runs[0].bold = True
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')


doc.add_heading('7. API Endpoints (Summary)', level=1)
for item in [
    'Auth: /api/auth/register, /api/auth/login, /api/auth/me',
    'Portfolio: /api/portfolios/:id',
    'Holdings: /api/portfolios/:id/holdings, /api/holdings/:id, /api/holdings/:id/price',
    'Live Quotes: /api/market/quote/:symbol, /api/portfolios/:id/holdings/refresh-prices',
    'Transactions: /api/portfolios/:id/transactions, /api/transactions/:id',
    'Chatbot: /api/chat/message'
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('8. Testing & Validation', level=1)
for item in [
    'Authentication tested for valid and invalid credentials.',
    'CRUD flows verified for holdings and transactions.',
    'Live quote and refresh endpoints validated against market API responses.',
    'Frontend build executed successfully after integration updates.',
    'Error handling improved using toast notifications for better UX.'
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('9. Results', level=1)
doc.add_paragraph(
    'The developed system successfully provides secure access, real-time portfolio tracking, '
    'automated metric calculations, and visual analytics. The integration of live-price refresh and '
    'AI assistance enhances usability and decision support for end users.'
)


doc.add_heading('10. Limitations', level=1)
for item in [
    'Market data quality depends on external API availability and symbol mapping.',
    'No advanced technical indicators or predictive analytics in current version.',
    'Chatbot accuracy depends on selected provider model and prompt context.'
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('11. Future Enhancements', level=1)
for item in [
    'Add watchlist and alert notifications for price thresholds.',
    'Add portfolio diversification and sector analytics.',
    'Implement role-based admin dashboard and audit logs.',
    'Integrate broker APIs for automatic transaction sync.',
    'Deploy with CI/CD and production monitoring.'
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('12. Conclusion', level=1)
doc.add_paragraph(
    'StockSense demonstrates a practical, scalable mini-project implementation of modern full-stack web development '
    'using MVC, secure authentication, cloud database integration, and data visualization. '
    'It meets core academic and functional objectives for a Stock Portfolio Tracker with real-world applicability.'
)

doc.save(output_path)
print(output_path)
